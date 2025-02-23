import shutil
from docx import Document
from docx.shared import Pt

def add_text_to_cell(cell, text):
    font = cell.paragraphs[0].runs[0].font.name
    font_size = cell.paragraphs[0].runs[0].font.size
    is_bold = cell.paragraphs[0].runs[0].bold
    is_italic = cell.paragraphs[0].runs[0].italic
    is_underlined = cell.paragraphs[0].runs[0].underline

    run = cell.paragraphs[0].add_run(text)
    run.font.name = font
    run.font.size = font_size
    if is_bold: run.bold = True
    if is_italic: run.italic = True
    if is_underlined: run.underline = True

src_path = "Addition_Questions.docx"
target_path = "Addition_Questions_ans.docx"
shutil.copyfile(src_path, target_path)

doc = Document(target_path)
table = doc.tables[0]

for row in table.rows:
    for cell in row.cells:
        ans = ' ' + str(eval(cell.text.strip('=')))
        add_text_to_cell(cell, ans)

doc.save(target_path)